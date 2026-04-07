#html_todocx_tables.py
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def get_class_names(node):
    classes = node.get("class", [])
    if isinstance(classes, str):
        classes = classes.split()
    return classes or []


def merge_style_maps(*style_maps):
    merged = {}
    for style_map in style_maps:
        if style_map:
            merged.update(style_map)
    return merged


def extract_class_style_map(node, config):
    class_styles = config.get("class_styles", {})
    classes = get_class_names(node)

    merged = {}
    for class_name in classes:
        class_style = class_styles.get(class_name, {})
        if class_style:
            merged.update(class_style)

    return merged


def extract_table_style_map(table_node, parser):
    class_style_map = extract_class_style_map(table_node, parser.config)

    inline_style_map = {}
    if parser.config.get("respect_inline_table_styles", True):
        style_str = table_node.get("style", "")
        inline_style_map = parser.parse_dict_string(style_str) if style_str else {}

    return merge_style_maps(class_style_map, inline_style_map)


def extract_cell_style_map(cell_node, parser):
    class_style_map = extract_class_style_map(cell_node, parser.config)

    inline_style_map = {}
    if parser.config.get("respect_inline_cell_styles", True):
        style_str = cell_node.get("style", "")
        inline_style_map = parser.parse_dict_string(style_str) if style_str else {}

    return merge_style_maps(class_style_map, inline_style_map)


def css_color_to_hex(parser, color_str):
    if not color_str:
        return None

    color_str = color_str.strip()

    if color_str.startswith("#"):
        color = color_str.lstrip("#")
        if len(color) == 3:
            color = "".join(ch * 2 for ch in color)
        return color.upper() if len(color) == 6 else None

    rgb = parser._parse_color(color_str)
    if rgb:
        return "".join(f"{x:02X}" for x in rgb)

    return None


def parse_css_border(border_value):
    if not border_value:
        return None

    parts = border_value.strip().split()
    if not parts:
        return None

    size = 4
    style = "single"
    color = "auto"

    css_to_docx_style = {
        "solid": "single",
        "dashed": "dashed",
        "dotted": "dotted",
        "double": "double",
    }

    for part in parts:
        part = part.strip().lower()

        if part.endswith("px"):
            try:
                px = float(part.replace("px", ""))
                size = max(2, int(px * 4))
            except ValueError:
                pass
        elif part in css_to_docx_style:
            style = css_to_docx_style[part]
        elif part.startswith("#") or part.startswith("rgb"):
            parsed = css_color_to_hex_from_str(part)
            if parsed:
                color = parsed

    return {"style": style, "size": size, "color": color}


def css_color_to_hex_from_str(color_str):
    color_str = color_str.strip()

    if color_str.startswith("#"):
        color = color_str.lstrip("#")
        if len(color) == 3:
            color = "".join(ch * 2 for ch in color)
        return color.upper() if len(color) == 6 else None

    return None


def build_border_dict_from_style_map(style_map, default_border=None):
    borders = {}

    if default_border:
        borders = {
            "top": default_border.copy(),
            "right": default_border.copy(),
            "bottom": default_border.copy(),
            "left": default_border.copy(),
        }

    if "border" in style_map:
        parsed = parse_css_border(style_map["border"])
        if parsed:
            borders = {
                "top": parsed.copy(),
                "right": parsed.copy(),
                "bottom": parsed.copy(),
                "left": parsed.copy(),
            }

    for side in ("top", "right", "bottom", "left"):
        key = f"border-{side}"
        if key in style_map:
            parsed = parse_css_border(style_map[key])
            if parsed:
                borders[side] = parsed

    return borders


def set_cell_shading(cell, fill):
    if not fill:
        return

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, borders):
    if not borders:
        return

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "right", "bottom", "left"):
        border = borders.get(edge)
        if not border:
            continue

        edge_tag = tc_borders.find(qn(f"w:{edge}"))
        if edge_tag is None:
            edge_tag = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_tag)

        edge_tag.set(qn("w:val"), border.get("style", "single"))
        edge_tag.set(qn("w:sz"), str(border.get("size", 4)))
        edge_tag.set(qn("w:color"), border.get("color", "auto"))
        edge_tag.set(qn("w:space"), "0")


def set_table_borders(table, borders):
    if not borders:
        return

    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge in ("top", "right", "bottom", "left", "insideH", "insideV"):
        border = borders.get(edge)
        if not border:
            continue

        edge_tag = tbl_borders.find(qn(f"w:{edge}"))
        if edge_tag is None:
            edge_tag = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_tag)

        edge_tag.set(qn("w:val"), border.get("style", "single"))
        edge_tag.set(qn("w:sz"), str(border.get("size", 4)))
        edge_tag.set(qn("w:color"), border.get("color", "auto"))
        edge_tag.set(qn("w:space"), "0")


def handle_table(parser):
    if not hasattr(parser, 'tables') or parser.table_no >= len(parser.tables):
        return

    soup_table = parser.tables[parser.table_no]
    row_tags = soup_table.find_all("tr")

    if not row_tags:
        parser.instances_to_skip = len(soup_table.find_all("table"))
        parser.skip = True
        parser.skip_tag = "table"
        return

    grid = []
    max_cols = 0

    for row in row_tags:
        grid_row = []
        for cell in row.find_all(["td", "th"], recursive=False):
            try:
                colspan = max(1, int(cell.get("colspan", 1)))
                rowspan = max(1, int(cell.get("rowspan", 1)))
            except (ValueError, TypeError):
                colspan, rowspan = 1, 1

            grid_row.append({
                "node": cell,
                "colspan": colspan,
                "rowspan": rowspan,
            })

        max_cols = max(max_cols, sum(c["colspan"] for c in grid_row))
        grid.append(grid_row)

    rows = len(grid)
    cols = max_cols
    span_map = [[None for _ in range(cols)] for _ in range(rows)]

    for r, row in enumerate(grid):
        c = 0
        for cell in row:
            while c < cols and span_map[r][c] is not None:
                c += 1

            for dr in range(cell["rowspan"]):
                for dc in range(cell["colspan"]):
                    rr = r + dr
                    cc = c + dc
                    if rr < rows and cc < cols:
                        span_map[rr][cc] = {
                            "cell": cell,
                            "anchor_r": r,
                            "anchor_c": c,
                            "is_anchor": dr == 0 and dc == 0,
                        }

            c += cell["colspan"]

    parser.table = parser.doc.add_table(rows=rows, cols=cols)

    table_style = parser.config.get("table_style")
    if table_style:
        try:
            parser.table.style = table_style
        except Exception:
            pass

    table_style_map = extract_table_style_map(soup_table, parser)
    default_table_border = parser.config.get("default_table_border")

    table_borders = build_border_dict_from_style_map(
        table_style_map,
        default_border=default_table_border
    )

    if table_borders:
        table_border_all = None
        if "border" in table_style_map:
            table_border_all = parse_css_border(table_style_map["border"])

        if "insideH" not in table_borders:
            if table_border_all:
                table_borders["insideH"] = table_border_all.copy()
            elif default_table_border:
                table_borders["insideH"] = default_table_border.copy()

        if "insideV" not in table_borders:
            if table_border_all:
                table_borders["insideV"] = table_border_all.copy()
            elif default_table_border:
                table_borders["insideV"] = default_table_border.copy()

        set_table_borders(parser.table, table_borders)

    for r in range(rows):
        for c in range(cols):
            meta = span_map[r][c]
            if not meta or not meta["is_anchor"]:
                continue

            cell_info = meta["cell"]
            docx_cell = parser.table.cell(r, c)

            html_inside = parser.get_cell_html(cell_info["node"])
            if cell_info["node"].name == "th":
                html_inside = f"<b>{html_inside}</b>"

            child = parser.__class__(parser.config)
            child.copy_settings_from(parser)
            child.add_html_to_cell(html_inside, docx_cell)

            cell_style_map = extract_cell_style_map(cell_info["node"], parser)

            bg_color = None
            if "background-color" in cell_style_map:
                bg_color = css_color_to_hex(parser, cell_style_map["background-color"])
            elif parser.config.get("default_cell_shading"):
                bg_color = parser.config.get("default_cell_shading")

            if bg_color:
                set_cell_shading(docx_cell, bg_color)

            default_cell_border = parser.config.get("default_cell_border")
            cell_borders = build_border_dict_from_style_map(
                cell_style_map,
                default_border=default_cell_border
            )

            if cell_borders:
                set_cell_borders(docx_cell, cell_borders)

            rs = cell_info["rowspan"]
            cs = cell_info["colspan"]

            if rs > 1 or cs > 1:
                end_r = min(r + rs - 1, rows - 1)
                end_c = min(c + cs - 1, cols - 1)
                try:
                    docx_cell.merge(parser.table.cell(end_r, end_c))
                except Exception:
                    pass

    parser.instances_to_skip = len(soup_table.find_all("table"))
    parser.skip = True
    parser.skip_tag = "table"
    parser.table = None