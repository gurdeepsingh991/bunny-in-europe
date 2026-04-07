import re
from html.parser import HTMLParser

import docx
import docx.table
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from .config import build_config
from .html_cleanup import preprocess_html
from .run_styles import (
    add_styles_to_run,
    apply_default_alignment,
    apply_default_run_font,
    apply_heading_run_style,
    apply_semantic_run_styles,
    parse_color,
    set_run_font,
)
from .utils import (
    INDENT,
    LIST_INDENT,
    MAX_INDENT,
    delete_paragraph,
    fetch_image,
    get_filename_from_url,
    is_url,
    remove_last_occurrence,
    remove_whitespace,
)

try:
    from .html_to_docx_tables import handle_table
except ImportError:
    from html_to_docx_tables import handle_table


DEFAULT_TABLE_STYLE = None
DEFAULT_PARAGRAPH_STYLE = None

LIST_STYLES = {
    "LIST_BULLET": "List Bullet",
    "LIST_NUMBER": "List Number",
}


class HtmlToDocx(HTMLParser):
    def __init__(self, config=None):
        super().__init__()
        self.options = {
            "fix-html": True,
            "images": True,
            "tables": True,
            "styles": True,
        }
        self.table_row_selectors = [
            "table > tr",
            "table > thead > tr",
            "table > tbody > tr",
            "table > tfoot > tr",
        ]
        self.table_style = DEFAULT_TABLE_STYLE
        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE
        self.config = build_config(config)

    def set_initial_attrs(self, document=None):
        self.tags = {"span": [], "list": []}
        self.doc = document if document else Document()
        self.bs = self.options["fix-html"]
        self.document = self.doc
        self.include_tables = True
        self.include_images = self.options["images"]
        self.include_styles = self.options["styles"]
        self.paragraph = None
        self.run = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style
        self.config = other.config.copy()

    def get_cell_html(self, soup):
        return " ".join(str(i) for i in soup.contents)

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as exc:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from exc

    def parse_dict_string(self, string, separator=";"):
        result = {}
        for item in string.split(separator):
            if ":" in item:
                key, value = item.split(":", 1)
                result[key.strip().lower()] = value.strip()
        return result

    def add_styles_to_paragraph(self, style):
        if "text-align" in style:
            align = style["text-align"]
            if align == "center":
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "justify":
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if "margin-left" in style:
            margin = style["margin-left"]
            units = re.sub(r"[0-9]+", "", margin)
            margin_val = int(float(re.sub(r"[a-z]+", "", margin)))
            if units == "px":
                self.paragraph.paragraph_format.left_indent = Inches(
                    min(margin_val // 10 * INDENT, MAX_INDENT)
                )

    def handle_li(self):
        list_depth = len(self.tags["list"])
        list_type = self.tags["list"][-1] if list_depth else "ul"
        list_style = (
            LIST_STYLES["LIST_NUMBER"] if list_type == "ol" else LIST_STYLES["LIST_BULLET"]
        )

        self.paragraph = self.doc.add_paragraph(style=list_style)
        self.paragraph.paragraph_format.left_indent = Inches(
            min(list_depth * LIST_INDENT, MAX_INDENT)
        )
        self.paragraph.paragraph_format.line_spacing = 1

    def add_image_to_cell(self, cell, image):
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image)

    def handle_img(self, current_attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = "img"
            return

        src = current_attrs.get("src", "")
        src_is_url = is_url(src)

        image = fetch_image(src) if src_is_url else src

        if image:
            try:
                if isinstance(self.doc, docx.document.Document):
                    self.doc.add_picture(image)
                else:
                    self.add_image_to_cell(self.doc, image)
            except FileNotFoundError:
                image = None

        if not image:
            placeholder = (
                f"<image: {src}>"
                if src_is_url
                else f"<image: {get_filename_from_url(src)}>"
            )
            self.doc.add_paragraph(placeholder)

    def handle_link(self, href, text):
        rel_id = self.paragraph.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )

        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(docx.oxml.shared.qn("r:id"), rel_id)

        subrun = self.paragraph.add_run(text)
        set_run_font(
            subrun,
            font_name=self.config.get("fonts", {}).get("default", {}).get("name"),
            font_size=self.config.get("fonts", {}).get("default", {}).get("size"),
            font_color=self.config.get("fonts", {}).get("default", {}).get("color"),
        )
        subrun.font.underline = True
        subrun.font.color.rgb = RGBColor(0x00, 0x00, 0xEE)
        hyperlink.append(subrun._r)
        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return

        if tag == "head":
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return
        if tag == "body":
            return

        current_attrs = dict(attrs)

        if tag == "span":
            self.tags["span"].append(current_attrs)
            return
        if tag in ["ol", "ul"]:
            self.tags["list"].append(tag)
            return
        if tag == "br":
            if self.run:
                self.run.add_break()
            return

        self.tags[tag] = current_attrs

        if tag in ["p", "pre"]:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()
        elif tag == "li":
            self.handle_li()
        elif tag == "hr":
            self._add_hr()
        elif re.match(r"h[1-9]", tag):
            self.paragraph = self.doc.add_paragraph()
        elif tag == "img":
            self.handle_img(current_attrs)
            return
        elif tag == "table":
            handle_table(self)
            return

        if self.include_styles and "style" in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs["style"])
            self.add_styles_to_paragraph(style)

    def _add_hr(self):
        self.paragraph = self.doc.add_paragraph()
        p_pr = self.paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        p_pr.insert_element_before(
            p_bdr,
            "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
            "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
            "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
            "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
            "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
            "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
            "w:pPrChange"
        )
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)

    def handle_endtag(self, tag):
        if self.skip:
            if tag != self.skip_tag:
                return

            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return

            self.skip = False
            self.skip_tag = None
            self.paragraph = None

        if tag == "span":
            if self.tags["span"]:
                self.tags["span"].pop()
                return
        elif tag in ["ol", "ul"]:
            remove_last_occurrence(self.tags["list"], tag)
            return
        elif tag == "table":
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None

        if tag in self.tags:
            self.tags.pop(tag)

    def handle_data(self, data):
        if self.skip:
            return

        if "pre" not in self.tags:
            data = remove_whitespace(data, True, True)

        if not data:
            return

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        link = self.tags.get("a")
        if link:
            self.handle_link(link["href"], data)
            return

        self.run = self.paragraph.add_run(data)
        in_table = not isinstance(self.doc, docx.document.Document)

        apply_default_run_font(self.run, self.config, in_table=in_table)
        apply_heading_run_style(self.run, self.tags, self.config)
        apply_semantic_run_styles(self.run, self.tags)
        apply_default_alignment(self.paragraph, self.paragraph.text, in_table=in_table)

    def ignore_nested_tables(self, tables_soup):
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all("table"))
        return new_tables

    def get_tables(self):
        if not hasattr(self, "soup"):
            self.include_tables = False
            return

        self.tables = self.ignore_nested_tables(self.soup.find_all("table"))
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, "html.parser")
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError("First argument needs to be a string")
        if not isinstance(document, docx.document.Document):
            raise ValueError("Second argument needs to be a Document")

        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError("Second argument needs to be a Cell")

        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)

        self.set_initial_attrs(cell)
        self.run_process(html)
        if not self.doc.paragraphs:
            self.doc.add_paragraph("")

    def parse_html_file(self, filename_html, filename_docx=None):
        with open(filename_html, "r", encoding="utf-8") as infile:
            html = infile.read()

        self.set_initial_attrs()
        self.run_process(html)

        if not filename_docx:
            import os
            path, filename = os.path.split(filename_html)
            filename_docx = f"{path}/new_docx_file_{filename}"

        self.doc.save(f"{filename_docx}.docx")

    def parse_html_string(self, html):
        html = preprocess_html(html, self.config)
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc